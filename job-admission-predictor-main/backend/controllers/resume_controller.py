import os
import re
from datetime import datetime
from functools import wraps
from flask import request, jsonify, current_app
from utils.auth_middleware import token_required

# Try to import file processing libraries
try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    PdfReader = None
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    pdfplumber = None
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    Document = None
    HAS_DOCX = False

# Skill keywords by role
ROLE_KEYWORDS = {
    'Frontend Developer': {
        'required': ['React', 'JavaScript', 'HTML', 'CSS', 'TypeScript', 'Vue', 'Angular'],
        'preferred': ['Tailwind', 'Bootstrap', 'Redux', 'Webpack', 'npm', 'Responsive Design', 'Figma']
    },
    'Full Stack Developer': {
        'required': ['React', 'Node.js', 'JavaScript', 'Python', 'SQL', 'MongoDB', 'REST API'],
        'preferred': ['Docker', 'Git', 'AWS', 'Microservices', 'Authentication', 'Database Design']
    },
    'ML Intern': {
        'required': ['Python', 'Machine Learning', 'TensorFlow', 'scikit-learn', 'Pandas', 'NumPy'],
        'preferred': ['Deep Learning', 'NLP', 'Computer Vision', 'Keras', 'PyTorch', 'Statistics']
    },
    'Data Analyst': {
        'required': ['SQL', 'Python', 'Data Analysis', 'Excel', 'Tableau', 'Power BI', 'Statistics'],
        'preferred': ['Pandas', 'Matplotlib', 'Visualization', 'A/B Testing', 'R', 'BigQuery']
    }
}

def normalize_extracted_text(text_parts):
    """Normalize extracted text by removing empty fragments and extra whitespace."""
    cleaned_parts = []

    for part in text_parts:
        if not part:
            continue

        normalized = "\n".join(line.rstrip() for line in str(part).splitlines())
        normalized = normalized.strip()
        if normalized:
            cleaned_parts.append(normalized)

    return "\n\n".join(cleaned_parts).strip()

def extract_text_from_pdf(file_path):
    """Extract text from PDF using PyPDF2 first, then pdfplumber as fallback."""
    extraction_errors = []

    if HAS_PYPDF2:
        try:
            reader = PdfReader(file_path)
            page_text_parts = []

            for page in reader.pages:
                page_text_parts.append(page.extract_text() or "")

            text = normalize_extracted_text(page_text_parts)
            if text:
                return text, None
        except Exception as e:
            extraction_errors.append(f"PyPDF2 failed: {e}")
    else:
        extraction_errors.append("PyPDF2 is not installed")

    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(file_path) as pdf:
                page_text_parts = []

                for page in pdf.pages:
                    page_text_parts.append(page.extract_text() or "")

            text = normalize_extracted_text(page_text_parts)
            if text:
                return text, None
        except Exception as e:
            extraction_errors.append(f"pdfplumber failed: {e}")
    else:
        extraction_errors.append("pdfplumber is not installed")

    return None, "; ".join(extraction_errors)

def extract_text_from_docx(file_path):
    """Extract text from DOCX paragraphs and tables."""
    if not HAS_DOCX:
        return None, "python-docx is not installed"

    try:
        doc = Document(file_path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs]

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        text = normalize_extracted_text(text_parts)
        if text:
            return text, None

        return None, "The DOCX file was read successfully but no text content was found"
    except Exception as e:
        return None, f"python-docx failed: {e}"

def extract_resume_text(file_path):
    """Extract text from a resume file and return text plus extraction status."""
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == '.pdf':
        text, error = extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        text, error = extract_text_from_docx(file_path)
    else:
        return {
            'text': None,
            'error': 'Unsupported file type',
            'method': None
        }

    return {
        'text': text,
        'error': error,
        'method': file_ext.lstrip('.')
    }

def extract_contact_info(text):
    """Extract contact information from resume text"""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    email = re.search(email_pattern, text)
    phone = re.search(phone_pattern, text)
    
    return {
        'has_email': bool(email),
        'has_phone': bool(phone),
        'email': email.group(0) if email else None,
        'phone': phone.group(0) if phone else None
    }

def detect_sections(text):
    """Detect presence of key resume sections"""
    text_lower = text.lower()
    
    sections = {
        'education': bool(re.search(r'\b(education|degree|bachelor|master|b\.?tech|b\.?e\.?|m\.?tech|m\.?e\.?)\b', text_lower)),
        'experience': bool(re.search(r'\b(experience|work|employment|job|position)\b', text_lower)),
        'projects': bool(re.search(r'\b(project|projects|built|developed|created)\b', text_lower)),
        'skills': bool(re.search(r'\b(skills|technical|programming|languages|tools|expertise)\b', text_lower))
    }
    
    return sections

def detect_skills(text, role='Full Stack Developer'):
    """Detect skills present in resume"""
    # Common technical skills to look for
    all_skills = {
        'Languages': ['Python', 'JavaScript', 'Java', 'C++', 'C#', 'Ruby', 'Go', 'Rust', 'PHP', 'Swift'],
        'Frontend': ['React', 'Vue', 'Angular', 'HTML', 'CSS', 'Tailwind', 'Bootstrap', 'Webpack'],
        'Backend': ['Node.js', 'Django', 'Flask', 'Spring', 'Express', 'FastAPI', 'ASP.NET'],
        'Databases': ['MongoDB', 'PostgreSQL', 'MySQL', 'SQL', 'Redis', 'Firebase', 'DynamoDB'],
        'ML/Data': ['TensorFlow', 'PyTorch', 'scikit-learn', 'Pandas', 'NumPy', 'Keras', 'Tableau', 'Power BI'],
        'DevOps': ['Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'CI/CD', 'Jenkins', 'Git'],
        'Other': ['REST API', 'Microservices', 'GraphQL', 'Agile', 'Linux', 'Authentication']
    }
    
    detected = []
    for category, skills in all_skills.items():
        for skill in skills:
            if re.search(rf'\b{skill}\b', text, re.IGNORECASE):
                detected.append(skill)
    
    # Get role-specific missing skills
    role_keywords = ROLE_KEYWORDS.get(role, {})
    required_skills = role_keywords.get('required', [])
    preferred_skills = role_keywords.get('preferred', [])
    
    missing_required = [s for s in required_skills if s not in detected]
    missing_preferred = [s for s in preferred_skills if s not in detected]
    
    return {
        'detected': list(set(detected)),
        'missing_required': missing_required[:5],  # Top 5 missing required
        'missing_preferred': missing_preferred[:5]   # Top 5 missing preferred
    }

def calculate_ats_score(resume_data, role='Full Stack Developer'):
    """Calculate estimated ATS score (0-100)"""
    score = 0
    breakdown = {}
    
    # Contact Info (15 points)
    contact_score = 0
    if resume_data['contact_info']['has_email']:
        contact_score += 7
    if resume_data['contact_info']['has_phone']:
        contact_score += 8
    score += contact_score
    breakdown['contact_info'] = contact_score
    
    # Sections (25 points)
    sections_score = 0
    sections = resume_data['sections']
    if sections.get('education'): sections_score += 8
    if sections.get('experience'): sections_score += 8
    if sections.get('projects'): sections_score += 5
    if sections.get('skills'): sections_score += 4
    score += sections_score
    breakdown['sections'] = sections_score
    
    # Skills Detection (20 points)
    skills_data = resume_data['skills']
    detected_count = len(skills_data['detected'])
    if detected_count >= 8: 
        skills_score = 20
    elif detected_count >= 5:
        skills_score = 15
    elif detected_count >= 3:
        skills_score = 10
    else:
        skills_score = 5
    score += skills_score
    breakdown['skills_detection'] = skills_score
    
    # Role Relevance (25 points)
    role_score = 0
    missing_required = len(skills_data['missing_required'])
    missing_preferred = len(skills_data['missing_preferred'])
    
    if missing_required == 0:
        role_score = 25
    elif missing_required <= 2:
        role_score = 20
    elif missing_required <= 4:
        role_score = 15
    else:
        role_score = 8
    
    # Adjust for missing preferred skills
    role_score -= (missing_preferred * 0.5)
    role_score = max(0, min(25, role_score))
    
    score += role_score
    breakdown['role_relevance'] = role_score
    
    # Formatting (15 points) - simplified heuristic
    text = resume_data['text']
    formatting_score = 10  # Base score
    
    # Check if text is well-formatted (has good content length)
    if len(text) > 800:
        formatting_score += 3
    if text.count('\n') > 10:
        formatting_score += 2
    
    score += formatting_score
    breakdown['formatting'] = formatting_score
    
    # Ensure score is between 0-100
    score = max(0, min(100, score))
    
    return {
        'estimated_ats_score': round(score, 1),
        'breakdown': breakdown
    }

def generate_suggestions(resume_data, ats_data, role='Full Stack Developer'):
    """Generate improvement suggestions"""
    suggestions = []
    score = ats_data['estimated_ats_score']
    
    # Contact info suggestions
    if not resume_data['contact_info']['has_email']:
        suggestions.append("Add your email address at the top of your resume")
    if not resume_data['contact_info']['has_phone']:
        suggestions.append("Include your phone number for recruiters to contact you")
    
    # Section suggestions
    sections = resume_data['sections']
    if not sections.get('education'):
        suggestions.append("Add an Education section with degree and institution details")
    if not sections.get('experience'):
        suggestions.append("Include a Work Experience section with your roles and achievements")
    if not sections.get('projects'):
        suggestions.append("Highlight 2-3 of your best projects with clear descriptions and technologies used")
    if not sections.get('skills'):
        suggestions.append("Add a dedicated Skills section organized by category (Languages, Frameworks, Tools)")
    
    # Skills suggestions
    missing_required = resume_data['skills']['missing_required']
    if missing_required:
        top_missing = missing_required[0]
        suggestions.append(f"Consider adding {top_missing} - it's a key requirement for {role} roles")
    
    # General tips
    if score < 50:
        suggestions.append("Your resume could use significant updates - focus on adding required skills and completing all sections")
    elif score < 70:
        suggestions.append("Add more technical keywords relevant to your target role to improve ATS compatibility")
    elif score < 85:
        suggestions.append("Consider quantifying your achievements (e.g., 'Improved performance by 30%')")
    else:
        suggestions.append("Your resume has excellent structure and keyword coverage!")
    
    return suggestions[:5]  # Return top 5 suggestions

@token_required
def analyze_resume(current_user):
    """Main endpoint to analyze resume"""
    try:
        # Check if file is present
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        role = request.form.get('role', 'Full Stack Developer')
        
        # Validate file
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        allowed_extensions = {'.pdf', '.docx'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            return jsonify({'error': 'Only PDF and DOCX files are supported'}), 400
        
        # Create uploads directory if it doesn't exist
        uploads_dir = os.path.join(os.path.dirname(__file__), '..', 'uploads')
        if not os.path.exists(uploads_dir):
            os.makedirs(uploads_dir, exist_ok=True)
        
        # Save file temporarily with cross-platform path handling
        filename = f'resume_{int(datetime.now().timestamp())}{file_ext}'
        temp_path = os.path.join(uploads_dir, filename)
        file.save(temp_path)
        
        # Extract text
        extraction_result = extract_resume_text(temp_path)
        resume_text = extraction_result['text']

        if not resume_text:
            os.remove(temp_path)
            extraction_error = extraction_result['error'] or 'No readable text found in the uploaded file'
            return jsonify({'error': extraction_error}), 400
        
        # Analyze resume
        contact_info = extract_contact_info(resume_text)
        sections = detect_sections(resume_text)
        skills = detect_skills(resume_text, role)
        
        resume_data = {
            'text': resume_text,
            'contact_info': contact_info,
            'sections': sections,
            'skills': skills
        }
        
        # Calculate ATS score
        ats_data = calculate_ats_score(resume_data, role)
        
        # Generate suggestions
        suggestions = generate_suggestions(resume_data, ats_data, role)
        
        # Generate summary
        score = ats_data['estimated_ats_score']
        if score >= 80:
            summary = "Excellent! Your resume is well-structured and has strong keyword coverage."
        elif score >= 60:
            summary = "Good! Your resume covers the basics, but adding more target keywords could improve it."
        else:
            summary = "Your resume needs more work. Focus on key skills and completing all sections."
        
        # Save analysis to database if available
        if current_app.db is not None:
            try:
                analysis = {
                    'user_id': current_user.get('_id'),
                    'email': current_user.get('email'),
                    'role': role,
                    'ats_score': ats_data['estimated_ats_score'],
                    'detected_skills': skills['detected'],
                    'missing_required_skills': skills['missing_required'],
                    'sections_present': sections,
                    'created_at': datetime.now()
                }
                current_app.db.resume_analyses.insert_one(analysis)
            except Exception as e:
                print(f"Error saving to DB: {e}")
        
        # Clean up
        os.remove(temp_path)
        
        return jsonify({
            'success': True,
            'estimated_ats_score': ats_data['estimated_ats_score'],
            'summary': summary,
            'detected_skills': skills['detected'][:10],
            'missing_required_skills': skills['missing_required'],
            'missing_preferred_skills': skills['missing_preferred'],
            'sections_present': sections,
            'contact_present': {
                'email': contact_info['has_email'],
                'phone': contact_info['has_phone']
            },
            'suggestions': suggestions,
            'breakdown': ats_data['breakdown'],
            'role_analyzed': role
        }), 200
    
    except Exception as e:
        return jsonify({'error': f'Error analyzing resume: {str(e)}'}), 500
